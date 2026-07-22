"""Document ingestion utilities shared by Lab 01.

Optional parsers are imported only when the matching file type is present, so a
Markdown-only workshop remains lightweight.  Install the extras in
``requirements.txt`` to enable PDF/DOCX/XLSX and image OCR ingestion.
"""
from __future__ import annotations

import re
import unicodedata
from pathlib import Path

SUPPORTED_SUFFIXES = {".md", ".txt", ".pdf", ".docx", ".xlsx", ".xls", ".png", ".jpg", ".jpeg", ".tif", ".tiff"}


def normalize_vietnamese(text: str) -> str:
    """Keep Unicode Vietnamese intact while removing invisible/noisy whitespace."""
    text = unicodedata.normalize("NFC", text).replace("\u00a0", " ")
    text = re.sub(r"[ \t]+", " ", text)
    return re.sub(r"\n{3,}", "\n\n", text).strip()


def extract_document(path: Path) -> list[dict]:
    """Return page/sheet-aware text units with source metadata.

    Every unit has ``text``, ``source`` and ``page``.  An empty/unsupported
    document fails loudly instead of silently producing a partial index.
    """
    suffix = path.suffix.lower()
    if suffix in {".md", ".txt"}:
        return [{"text": normalize_vietnamese(path.read_text(encoding="utf-8")), "source": path.name, "page": 1}]
    if suffix == ".pdf":
        import fitz  # PyMuPDF
        with fitz.open(path) as pdf:
            return [{"text": normalize_vietnamese(page.get_text("text")), "source": path.name, "page": i + 1} for i, page in enumerate(pdf)]
    if suffix == ".docx":
        from docx import Document
        doc = Document(path)
        text = "\n".join(p.text for p in doc.paragraphs)
        for table in doc.tables:
            text += "\n" + "\n".join(" | ".join(c.text for c in row.cells) for row in table.rows)
        return [{"text": normalize_vietnamese(text), "source": path.name, "page": 1}]
    if suffix in {".xlsx", ".xls"}:
        # openpyxl cannot read legacy binary .xls files. Pandas selects xlrd
        # for .xls and openpyxl for .xlsx.
        import pandas as pd
        units = []
        sheets = pd.read_excel(path, sheet_name=None, header=None)
        for sheet_name, frame in sheets.items():
            rows = [" | ".join("" if pd.isna(cell) else str(cell) for cell in row) for row in frame.itertuples(index=False, name=None)]
            units.append({"text": normalize_vietnamese("\n".join(rows)), "source": path.name, "page": sheet_name})
        return units
    if suffix in {".png", ".jpg", ".jpeg", ".tif", ".tiff"}:
        import pytesseract
        from PIL import Image
        return [{"text": normalize_vietnamese(pytesseract.image_to_string(Image.open(path), lang="vie+eng")), "source": path.name, "page": 1}]
    raise ValueError(f"Unsupported document type: {path.name}")


def document_files(data_dir: Path) -> list[Path]:
    files = (path for path in data_dir.iterdir() if path.is_file() and path.suffix.lower() in SUPPORTED_SUFFIXES and path.name.lower() != "readme.md")
    return sorted(files, key=lambda path: path.name.casefold())
